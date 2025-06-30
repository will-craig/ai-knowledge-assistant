from http.client import HTTPException
from contracts.search_query import SearchQuery
from fastapi import FastAPI, File, UploadFile

from ingest import extract_text_from_pdf
from docx import Document as DocxDocument

from llm_pipeline import get_embedding
from vectorstore.db import Document, SessionLocal, init_db, search_similar_documents

import os

app = FastAPI()

from fastapi.middleware.cors import CORSMiddleware

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # your frontend origin
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

init_db()
UPLOAD_DIR = "uploaded_pdfs"
os.makedirs(UPLOAD_DIR, exist_ok=True)

@app.post("/upload")
async def upload_pdf(file: UploadFile = File(...)):
    contents = await file.read()

    text = extract_text_from_upload(file, contents)
    
    # only take first 500 characters for embedding as limited on token count
    embedding = get_embedding(text[500])

    # Store in DB
    db = SessionLocal()
    doc = Document(filename=file.filename, content=text, embedding=embedding)
    db.add(doc)
    db.commit()
    db.refresh(doc)
    db.close()

    return {
        "filename": file.filename,
        "text_preview": text[:1000],
        "embedding_dim": len(embedding)
    }

@app.post("/search")
async def search_docs(search: SearchQuery):
    embedding = get_embedding(search.query)
    session = SessionLocal()
    try:
        results = search_similar_documents(session, embedding, search.top_k)
        return [
            {
                "filename": row.filename,
                "score": row.score,
                "preview": row.content[:300]
            }
            for row in results
        ]
    finally:
        session.close()

def extract_text_from_upload(file, contents):
    if file.filename.endswith(".pdf"):
        with open("temp.pdf", "wb") as f:
            f.write(contents)
        text = extract_text_from_pdf("temp.pdf")
        os.remove("temp.pdf")

    elif file.filename.endswith(".docx"):
        with open("temp.docx", "wb") as f:
            f.write(contents)
        text = extract_text_from_docx("temp.docx")
        os.remove("temp.docx")
    
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    return text

def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return "\n".join([p.text for p in doc.paragraphs])